#!/usr/bin/env python3
"""
Batch DOCX to MD Converter and Metadata Enricher

모든 DOCX 파일을 MD 로 변환한 후 JSONL metadata 를 일괄 업데이트합니다.

Usage:
    python functions/tools/batch_convert_and_enrich.py
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def extract_page_from_heading_text(text: str) -> tuple[str, Optional[int]]:
    """Extract page number from heading text if present."""
    match = re.match(r'^(.+?)\s+(\d+)$', text.strip())
    if match:
        title = match.group(1).strip()
        page_str = match.group(2)
        page_num = int(page_str)
        if 1 <= page_num <= 10000:
            return title, page_num
    return text, None


def convert_docx_to_md_with_pages(docx_path: Path) -> str:
    """Convert DOCX to Markdown with page number tracking."""
    doc = Document(docx_path)
    md_lines: list[str] = []
    last_known_page = 1
    
    page_number_pattern = re.compile(r'^\s*(\d+)\s*$')
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        if page_number_pattern.match(text):
            continue
        
        style_name = para.style.name if para.style else ""
        
        heading_match = re.match(r"^Heading\s*(\d+)", style_name, re.IGNORECASE)
        if heading_match:
            level = int(heading_match.group(1))
            prefix = "#" * level
            clean_text, page_num = extract_page_from_heading_text(text)
            if page_num:
                last_known_page = page_num
                md_lines.append(f"{prefix} {clean_text}\t{page_num}")
            else:
                md_lines.append(f"{prefix} {clean_text}\t{last_known_page}")
            continue
        
        section_match = re.match(r"^(\d+(\.\d+)+)\s+(.+)$", text)
        if section_match and len(section_match.group(1).split(".")) <= 6:
            section_num = section_match.group(1)
            section_title = section_match.group(3)
            level = min(len(section_num.split(".")), 6)
            prefix = "#" * level
            clean_title, page_num = extract_page_from_heading_text(section_title)
            if page_num:
                last_known_page = page_num
                md_lines.append(f"{prefix} {clean_title}\t{page_num}")
            else:
                md_lines.append(f"{prefix} {clean_title}\t{last_known_page}")
            continue
        
        if any(run.bold for run in para.runs if run.text.strip()):
            bold_text = "".join(run.text for run in para.runs if run.bold)
            if bold_text.strip() and len(bold_text.strip()) < 100:
                clean_text, page_num = extract_page_from_heading_text(text)
                if page_num:
                    last_known_page = page_num
                    md_lines.append(f"## {clean_text}\t{page_num}")
                else:
                    md_lines.append(f"## {clean_text}\t{last_known_page}")
                continue
        
        parts = text.split('\t')
        if len(parts) >= 2:
            last_part = parts[-1]
            if last_part.isdigit() and 1 <= int(last_part) <= 10000:
                md_lines.append(text)
                continue
        
        md_lines.append(text)
    
    for table in doc.tables:
        md_lines.append("")
        md_lines.append(f"## Table\t{last_known_page}")
        md_lines.append(convert_table_to_md(table))
        md_lines.append("")
    
    return "\n".join(md_lines)


def convert_table_to_md(table) -> str:
    """Convert DOCX table to Markdown."""
    md_lines: list[str] = []
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("|", "\\|").replace("\n", " ") for cell in row.cells]
        rows.append(cells)
    
    if not rows:
        return ""
    
    max_cols = max(len(row) for row in rows)
    for row in rows:
        while len(row) < max_cols:
            row.append("")
    
    for i, row in enumerate(rows):
        md_lines.append("| " + " | ".join(row) + " |")
        if i == 0:
            md_lines.append("| " + " | ".join("---" for _ in row) + " |")
    
    return "\n".join(md_lines)


def find_md_file(source_path: str, search_dirs: list[Path]) -> Optional[Path]:
    """Find MD file based on source_path."""
    rel_path = source_path.replace('data\\', 'data/').replace('data/', '')
    
    for search_dir in search_dirs:
        candidate = search_dir / rel_path
        if candidate.exists():
            return candidate
    
    if '3gpp_docs' in rel_path and 'raw' not in rel_path:
        rel_path = rel_path.replace('3gpp_docs', 'raw/3gpp_docs')
        for search_dir in search_dirs:
            candidate = search_dir / rel_path
            if candidate.exists():
                return candidate
    
    return None


def extract_page_info_from_md(md_path: Path) -> dict[int, int]:
    """Extract line number to page number mapping from MD file."""
    line_to_page: dict[int, int] = {}
    current_page = 1
    
    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        for line_num, line in enumerate(lines, 1):
            line = line.rstrip()
            
            heading_match = re.match(r'^(#+)\s+(.+?)\t(\d+)\s*$', line)
            if heading_match:
                page_num = int(heading_match.group(3))
                heading_text = heading_match.group(2)
                if 'page' not in heading_text.lower() and 'slide' not in heading_text.lower():
                    current_page = page_num
            
            elif not line.startswith('#') and '\t' in line:
                parts = line.split('\t')
                if len(parts) >= 2 and parts[-1].isdigit():
                    page_num = int(parts[-1])
                    if parts[0].replace('.', '').isdigit():
                        current_page = page_num
            
            line_to_page[line_num] = current_page
        
    except Exception as e:
        print(f"  [WARN] Error reading {md_path}: {e}")
    
    return line_to_page


def find_original_file(md_path: Path, search_dirs: list[Path]) -> tuple[Optional[str], str]:
    """Find original file (DOCX, PPTX, PDF) corresponding to MD file."""
    md_name = md_path.stem
    
    parent_dir = md_path.parent
    for ext in ['.docx', '.pptx', '.pdf']:
        candidate = parent_dir / f"{md_name}{ext}"
        if candidate.exists():
            return candidate.name, ext[1:].upper()
    
    for search_dir in search_dirs:
        for ext in ['.docx', '.pptx', '.pdf']:
            candidate = search_dir / f"{md_name}{ext}"
            if candidate.exists():
                return candidate.name, ext[1:].upper()
    
    if 'pptx' in md_name.lower():
        return f"{md_name}.pptx", "PPTX"
    elif 'pdf' in md_name.lower():
        return f"{md_name}.pdf", "PDF"
    else:
        return f"{md_name}.docx", "DOCX"


def enrich_chunk(chunk: dict, search_dirs: list[Path], md_cache: dict) -> dict:
    """Enrich chunk with source_file and page_number metadata."""
    source_path = chunk.get('source_path', '')
    if not source_path:
        return chunk
    
    if source_path not in md_cache:
        md_path = find_md_file(source_path, search_dirs)
        if md_path:
            source_file, file_format = find_original_file(md_path, search_dirs)
            line_to_page = extract_page_info_from_md(md_path)
            
            md_cache[source_path] = {
                'md_path': md_path,
                'source_file': source_file,
                'file_format': file_format,
                'line_to_page': line_to_page,
            }
        else:
            md_cache[source_path] = None
    
    file_info = md_cache.get(source_path)
    if not file_info:
        return chunk
    
    if 'metadata' not in chunk:
        chunk['metadata'] = {}
    
    if file_info['source_file']:
        chunk['metadata']['source_file'] = file_info['source_file']
    
    start_line = chunk.get('start_line', 1)
    page_num = file_info['line_to_page'].get(start_line)
    if page_num is not None:
        chunk['metadata']['page_number'] = page_num
    
    return chunk


def process_jsonl_file(jsonl_path: Path, search_dirs: list[Path], md_cache: dict) -> int:
    """Process a single JSONL file."""
    updated_count = 0
    updated_chunks = []
    
    try:
        with open(jsonl_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line:
                    chunk = json.loads(line)
                    chunk = enrich_chunk(chunk, search_dirs, md_cache)
                    updated_chunks.append(chunk)
                    updated_count += 1
        
        with open(jsonl_path, 'w', encoding='utf-8') as f:
            for chunk in updated_chunks:
                f.write(json.dumps(chunk, ensure_ascii=False) + '\n')
        
    except Exception as e:
        print(f"  [ERROR] {jsonl_path.name}: {e}")
        return 0
    
    return updated_count


def main():
    base_dir = Path('c:/Users/jy2601.kim/RAS-TestCaseReview')
    jsonl_dir = base_dir / 'data/processed/jsonl'
    
    search_dirs = [
        base_dir / 'data/raw',
        base_dir / 'data',
        base_dir / 'data/raw/3gpp_docs',
        base_dir / 'data/3gpp_docs',
    ]
    
    for subdir in ['algorithm_docs', 'hld_dld', 'feature_specs', 'issue_cases', 'legacy_tc']:
        subdir_path = base_dir / 'data' / subdir
        if subdir_path.exists():
            search_dirs.append(subdir_path)
    
    print("="*70)
    print("Batch DOCX to MD Converter and Metadata Enricher")
    print("="*70)
    
    # Step 1: Find all DOCX files without MD
    print("\n[Step 1] Finding DOCX files without MD...")
    docx_files_without_md = []
    
    for search_dir in search_dirs:
        if not search_dir.exists():
            continue
        for docx_file in search_dir.rglob("*.docx"):
            md_file = docx_file.with_suffix(".md")
            if not md_file.exists():
                docx_files_without_md.append(docx_file)
    
    print(f"  Found {len(docx_files_without_md)} DOCX files without MD")
    
    # Step 2: Convert DOCX to MD
    if docx_files_without_md:
        print(f"\n[Step 2] Converting {len(docx_files_without_md)} DOCX files to MD...")
        for i, docx_file in enumerate(docx_files_without_md, 1):
            md_file = docx_file.with_suffix(".md")
            try:
                print(f"  [{i}/{len(docx_files_without_md)}] {docx_file.name}...", end=" ")
                md_content = convert_docx_to_md_with_pages(docx_file)
                md_file.parent.mkdir(parents=True, exist_ok=True)
                md_file.write_text(md_content, encoding="utf-8")
                print(f"OK ({len(md_content)} chars)")
            except Exception as e:
                print(f"ERROR: {e}")
    else:
        print("  All DOCX files already have MD counterparts")
    
    # Step 3: Enrich JSONL metadata
    print(f"\n[Step 3] Enriching JSONL metadata...")
    jsonl_files = list(jsonl_dir.glob("*.jsonl"))
    print(f"  Found {len(jsonl_files)} JSONL files")
    
    md_cache: dict = {}
    total_updated = 0
    files_with_source = 0
    files_with_page = 0
    
    for i, jsonl_file in enumerate(jsonl_files, 1):
        print(f"  [{i}/{len(jsonl_files)}] {jsonl_file.name}...", end=" ")
        
        count = process_jsonl_file(jsonl_file, search_dirs, md_cache)
        total_updated += count
        
        if md_cache:
            first_info = next(iter(md_cache.values()))
            if first_info:
                if first_info.get('source_file'):
                    files_with_source += 1
                if first_info.get('line_to_page'):
                    files_with_page += 1
        
        print(f"OK ({count} chunks)")
    
    print()
    print("="*70)
    print(f"Total updated chunks: {total_updated}")
    print(f"Files with source_file metadata: {files_with_source}")
    print(f"Files with page_number mapping: {files_with_page}")
    print("="*70)


if __name__ == "__main__":
    main()
